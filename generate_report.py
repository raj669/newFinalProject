#!/usr/bin/env python3
"""
generate_report.py
==================
Converts PROJECT_REPORT.md into a submission-ready Word document
(PROJECT_REPORT.docx) following the prescribed format standards:

  - Page size : A4
  - Margins   : Top = 1 in, Bottom = 1 in, Right = 1 in, Left = 1.25 in
  - Font      : Times New Roman, 12 pt body
  - Spacing   : 1.5 line spacing, justified paragraphs
  - Headings  : 16 pt chapter titles (##), 14 pt sections (###),
                12 pt sub-sections (####) — all bold
  - Page nos  : Bottom-centre; Roman (i, ii …) for front matter,
                Arabic (1, 2 …) from Chapter 1 onwards
  - Tables    : Centred, caption centred above (bold 12 pt)
  - Code      : Courier New 9 pt, left-aligned block

Usage:
    python3 generate_report.py
    # → writes PROJECT_REPORT.docx in the same directory
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────────────────
REPO_DIR = Path(__file__).parent
MD_PATH  = REPO_DIR / "PROJECT_REPORT.md"
OUT_PATH = REPO_DIR / "PROJECT_REPORT.docx"

TNR     = "Times New Roman"
COURIER = "Courier New"

# ──────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ──────────────────────────────────────────────────────────────────────────────

def _xml_page_number(fmt: str = "decimal"):
    """Return XML run elements that produce a PAGE field in *fmt* format."""
    fld_begin   = OxmlElement("w:fldChar");  fld_begin.set(qn("w:fldCharType"),   "begin")
    instr        = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    if fmt == "lowerRoman":
        instr.text = " PAGE \\* lowerRoman "
    else:
        instr.text = " PAGE "
    fld_sep     = OxmlElement("w:fldChar");  fld_sep.set(qn("w:fldCharType"),     "separate")
    fld_end     = OxmlElement("w:fldChar");  fld_end.set(qn("w:fldCharType"),     "end")
    return fld_begin, instr, fld_sep, fld_end


def set_section_footer(section, fmt: str = "decimal"):
    """Place a centred page-number footer on *section*."""
    footer = section.footer
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.name = TNR
    run.font.size = Pt(12)
    for elem in _xml_page_number(fmt):
        run._r.append(elem)


def restart_numbering(section, start: int = 1):
    """Restart page numbering at *start* for *section*."""
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), str(start))
    # Remove any existing pgNumType
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    sectPr.append(pgNumType)


def enable_title_page(section):
    """Tell Word this section has a distinct first-page header/footer."""
    sectPr = section._sectPr
    titlePg = OxmlElement("w:titlePg")
    sectPr.append(titlePg)


def _set_page_layout(section):
    """A4, T=1 in, B=1 in, R=1 in, L=1.25 in."""
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.left_margin   = Inches(1.25)

# ──────────────────────────────────────────────────────────────────────────────
# Paragraph / run helpers
# ──────────────────────────────────────────────────────────────────────────────

def _para_fmt(para,
              align  = WD_ALIGN_PARAGRAPH.JUSTIFY,
              before = Pt(0),
              after  = Pt(6),
              ls     = 1.5):
    f = para.paragraph_format
    f.alignment         = align
    f.space_before      = before
    f.space_after       = after
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing      = ls


def _run_fmt(run, size=12, bold=False, italic=False, name=TNR):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic


def _apply_inline(para, text: str, size=12, base_bold=False, name=TNR):
    """
    Walk through *text* and emit bold / italic / code runs into *para*.
    Handles:  ***text***  **text**  *text*  `code`  [label](url)
    """
    i = 0
    n = len(text)
    while i < n:
        # *** bold-italic ***
        m = re.match(r'\*\*\*(.+?)\*\*\*', text[i:], re.DOTALL)
        if m:
            r = para.add_run(m.group(1))
            _run_fmt(r, size, bold=True, italic=True, name=name)
            i += len(m.group(0)); continue

        # ** bold **
        m = re.match(r'\*\*(.+?)\*\*', text[i:], re.DOTALL)
        if m:
            r = para.add_run(m.group(1))
            _run_fmt(r, size, bold=True, name=name)
            i += len(m.group(0)); continue

        # * italic *
        m = re.match(r'\*(.+?)\*', text[i:], re.DOTALL)
        if m:
            r = para.add_run(m.group(1))
            _run_fmt(r, size, bold=base_bold, italic=True, name=name)
            i += len(m.group(0)); continue

        # `inline code`
        m = re.match(r'`(.+?)`', text[i:], re.DOTALL)
        if m:
            r = para.add_run(m.group(1))
            _run_fmt(r, size=10, name=COURIER)
            i += len(m.group(0)); continue

        # [label](url) → show label only
        m = re.match(r'\[(.+?)\]\((.+?)\)', text[i:])
        if m:
            r = para.add_run(m.group(1))
            _run_fmt(r, size, bold=base_bold, name=name)
            i += len(m.group(0)); continue

        # plain text up to the next special character
        m = re.match(r'[^*`\[]+', text[i:])
        if m:
            r = para.add_run(m.group(0))
            _run_fmt(r, size, bold=base_bold, name=name)
            i += len(m.group(0)); continue

        # single unknown character
        r = para.add_run(text[i])
        _run_fmt(r, size, bold=base_bold, name=name)
        i += 1

# ──────────────────────────────────────────────────────────────────────────────
# Table helper
# ──────────────────────────────────────────────────────────────────────────────

def _is_separator_row(raw: str) -> bool:
    """True if the row is a markdown table separator  |---|:---:|---|."""
    cells = [c.strip() for c in raw.strip().strip("|").split("|")]
    return bool(cells) and all(re.match(r'^[-:]+$', c) for c in cells if c)


def _render_table(doc, md_rows: list):
    """Emit a docx table from collected markdown table rows."""
    data_rows = [r for r in md_rows if not _is_separator_row(r)]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    num_cols = max(len(r) for r in parsed)
    num_rows = len(parsed)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.style     = "Table Grid"
    tbl.alignment = 1  # WD_TABLE_ALIGNMENT.CENTER

    for ri, row_cells in enumerate(parsed):
        for ci in range(num_cols):
            cell_text = row_cells[ci] if ci < len(row_cells) else ""
            cell = tbl.rows[ri].cells[ci]
            # python-docx always creates one paragraph in a new cell
            para = cell.paragraphs[0]
            para.clear()
            _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                      before=Pt(2), after=Pt(2), ls=1.0)
            _apply_inline(para, cell_text, size=11,
                          base_bold=(ri == 0))

    # Blank spacer after table
    spacer = doc.add_paragraph()
    _para_fmt(spacer, after=Pt(6))

# ──────────────────────────────────────────────────────────────────────────────
# Main builder
# ──────────────────────────────────────────────────────────────────────────────

def build(lines: list) -> Document:
    doc = Document()

    # ── Default Normal style ─────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(12)

    # ── Initial section: title page + front matter ───────────────────────────
    s0 = doc.sections[0]
    _set_page_layout(s0)
    enable_title_page(s0)          # suppress page # on first (title) page
    set_section_footer(s0, "lowerRoman")   # Roman numerals for abstract/TOC
    restart_numbering(s0, start=1)

    # ── State ────────────────────────────────────────────────────────────────
    in_code_block  = False
    code_lines: list = []
    md_table_rows: list = []
    chapter_section_created = False
    # Track whether the first horizontal rule (title-page divider) has been seen.
    # Only the metadata lines before that divider should be centred.
    past_first_divider = False

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")

        # ── Code fence ───────────────────────────────────────────────────────
        if raw.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1; continue
            else:
                # close fence → emit block
                in_code_block = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                          before=Pt(4), after=Pt(8), ls=1.0)
                run = p.add_run(code_text)
                _run_fmt(run, size=9, name=COURIER)
                i += 1; continue

        if in_code_block:
            code_lines.append(raw)
            i += 1; continue

        # ── Markdown table row ───────────────────────────────────────────────
        if raw.startswith("|"):
            md_table_rows.append(raw)
            i += 1; continue
        else:
            if md_table_rows:
                _render_table(doc, md_table_rows)
                md_table_rows = []

        # ── Horizontal rule / blank ──────────────────────────────────────────
        stripped = raw.strip()
        if stripped in ("---", "***", "___") or not stripped:
            # Mark that we are past the title-page divider on the first `---`
            if stripped == "---" and not past_first_divider:
                past_first_divider = True
            i += 1; continue

        # ── H1: document title ───────────────────────────────────────────────
        if re.match(r'^# (?!#)', raw):
            text = raw[2:].strip()
            p = doc.add_paragraph()
            _para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                      before=Pt(72), after=Pt(30), ls=1.5)
            run = p.add_run(text)
            _run_fmt(run, size=18, bold=True)
            i += 1; continue

        # ── H2: chapter / top-level section ─────────────────────────────────
        if re.match(r'^## (?!#)', raw):
            text = raw[3:].strip()

            # ── Insert section break before Chapter 1 ───────────────────────
            if re.match(r'^Chapter 1\b', text) and not chapter_section_created:
                chapter_section_created = True
                s1 = doc.add_section(WD_SECTION.NEW_PAGE)
                _set_page_layout(s1)
                set_section_footer(s1, "decimal")
                restart_numbering(s1, start=1)

            p = doc.add_paragraph()
            _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                      before=Pt(18), after=Pt(12), ls=1.5)
            display = text.upper() if text.lower().startswith("chapter") else text
            run = p.add_run(display)
            _run_fmt(run, size=16, bold=True)
            i += 1; continue

        # ── H3: section ──────────────────────────────────────────────────────
        if re.match(r'^### (?!#)', raw):
            text = raw[4:].strip()
            p = doc.add_paragraph()
            _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                      before=Pt(14), after=Pt(8), ls=1.5)
            run = p.add_run(text)
            _run_fmt(run, size=14, bold=True)
            i += 1; continue

        # ── H4: sub-section ──────────────────────────────────────────────────
        if re.match(r'^#### (?!#)', raw):
            text = raw[5:].strip()
            p = doc.add_paragraph()
            _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                      before=Pt(10), after=Pt(6), ls=1.5)
            run = p.add_run(text)
            _run_fmt(run, size=12, bold=True)
            i += 1; continue

        # ── H5: sub-sub-section ──────────────────────────────────────────────
        if re.match(r'^##### (?!#)', raw):
            text = raw[6:].strip()
            p = doc.add_paragraph()
            _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                      before=Pt(8), after=Pt(4), ls=1.5)
            run = p.add_run(text)
            _run_fmt(run, size=12, bold=True, italic=True)
            i += 1; continue

        # ── Bullet list item ─────────────────────────────────────────────────
        if re.match(r'^[-*]\s', raw):
            text = raw[2:].strip()
            p = doc.add_paragraph()
            _para_fmt(p, before=Pt(0), after=Pt(3), ls=1.5)
            p.paragraph_format.left_indent = Inches(0.3)
            run0 = p.add_run("• ")
            _run_fmt(run0, size=12)
            _apply_inline(p, text)
            i += 1; continue

        # ── Indented bullet (two-space or tab prefix) ─────────────────────────
        if re.match(r'^  [-*]\s', raw):
            text = raw[4:].strip()
            p = doc.add_paragraph()
            _para_fmt(p, before=Pt(0), after=Pt(2), ls=1.5)
            p.paragraph_format.left_indent = Inches(0.55)
            run0 = p.add_run("◦ ")
            _run_fmt(run0, size=12)
            _apply_inline(p, text)
            i += 1; continue

        # ── Numbered list ─────────────────────────────────────────────────────
        m = re.match(r'^(\d+)\.\s+(.*)', raw)
        if m:
            num, text = m.group(1), m.group(2)
            p = doc.add_paragraph()
            _para_fmt(p, before=Pt(0), after=Pt(3), ls=1.5)
            p.paragraph_format.left_indent = Inches(0.3)
            run0 = p.add_run(f"{num}. ")
            _run_fmt(run0, size=12, bold=False)
            _apply_inline(p, text)
            i += 1; continue

        # ── Regular paragraph (including bold metadata cover lines) ──────────
        p = doc.add_paragraph()
        # Cover-page metadata: centred, no extra space.
        # Only apply centering to bold lines that appear before the first `---`
        # divider (the title-page block). Body paragraphs starting with **bold**
        # must not be centred.
        if not past_first_divider and stripped.startswith("**"):
            _para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                      before=Pt(0), after=Pt(4), ls=1.5)
        else:
            _para_fmt(p, before=Pt(0), after=Pt(6), ls=1.5)
        _apply_inline(p, stripped)
        i += 1

    # Flush any pending table rows
    if md_table_rows:
        _render_table(doc, md_table_rows)

    return doc

# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

def main():
    if not MD_PATH.exists():
        print(f"ERROR: {MD_PATH} not found.", file=sys.stderr)
        sys.exit(1)

    print(f"Reading  {MD_PATH.name} …")
    lines = MD_PATH.read_text(encoding="utf-8").splitlines(keepends=True)
    print(f"  {len(lines)} lines loaded.")

    print("Building Word document …")
    doc = build(lines)

    doc.save(OUT_PATH)
    print(f"Saved   → {OUT_PATH}")
    print("Done.")


if __name__ == "__main__":
    main()
